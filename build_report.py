import os
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

proj = r'C:\Тагир\Проект QC'
ss_dir = os.path.join(proj, 'screenshots')
out_path = os.path.join(proj, 'report_5P.docx')

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
# We use default font since Times New Roman may not be installed
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Helper
def add_centered_para(text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

add_centered_para(
    'Prototip avtomaticheskogo kontrolya kachestva\n'
    'rezultatov segmentatsii FreeSurfer', 18, True)

doc.add_paragraph('')

add_centered_para('Otchet po zadaniyu 5P kursa "Mini-VKR"', 14)

doc.add_paragraph('')
doc.add_paragraph('')

add_centered_para('Student: [FIO]', 12)
add_centered_para('Gruppa: [Gruppa]', 12)
add_centered_para('Data: iyun 2026 g.', 12)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Vvedenie', level=1)

doc.add_paragraph(
    'Nastoyashchiy otchet opisyvaet prototip sistemy avtomaticheskogo kontrolya '
    'kachestva (Quality Control, QC) rezultatov obrabotki MRT-dannykh golovnogo '
    'mozga payplaynom FreeSurfer. Prototip razrabotan v ramkakh zadaniya 5P kursa '
    '"Mini-VKR" i predstavlyaet soboy interaktivnoe veb-prilozhenie na baze Streamlit.'
)

doc.add_paragraph(
    'Aktualnost zadachi obuslovlena tem, chto standartnyy payplayn FreeSurfer '
    'ne vklyuchaet vstroennykh sredstv validatsii kachestva segmentatsii. '
    'Issledovatel vynuzhden vruchnuyu prosmotrivat desyatki i sotni srezov '
    'dlya kazhdogo subekta, chto trudoyomko i podverzheno chelovecheskim oshibkam. '
    'Avtomaticheskiy QC na osnove sravneniya s populyatsionnymi normami (Z-score) '
    'pozvolyaet bystro vyyavit problemnye regiony i sfokusirovat vnimanie '
    'eksperta na deystvitelno podozritelnykh strukturakh.'
)

doc.add_paragraph('Tsel raboty — sozdanie prototipa, kotoryy:')

goals = [
    'parsit vykhodnye fayly FreeSurfer (aseg.stats, lh/rh.aparc.stats) i izvlekaet '
    'izmerennye znacheniya obyomov podkorkovykh struktur i tolshchiny kory;',
    'sravnivaet izmereniya s normativnymi dannymi iz dvukh istochnikov — '
    'bazovye populyatsionnye normy (norms.csv) i nauchnye normy Brain Charts '
    '(Rutherford et al., Nature 2022);',
    'vychislyaet Z-score i klassifitsiruet kazhdyy region kak OK (|Z|<2), '
    'Check (2<=|Z|<3) ili Bad (|Z|>=3);',
    'predostavlyaet interaktivnyy dashbord s vizualizatsiey problemnykh zon '
    'na srezakh MRT, navigatsiey, filtratsiey i eksportom rezultatov.',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_page_break()

# ============================================================
# 2. FUNCTIONALITY
# ============================================================
doc.add_heading('2. Realizovannaya funktsionalnost', level=1)

doc.add_heading('2.1. Arkhitektura prototipa', level=2)
doc.add_paragraph(
    'Prototip sostoit iz 10 Python-moduley, organizovannykh po printsipu '
    'razdeleniya otvetstvennosti (Separation of Concerns):'
)

modules = [
    ('parser_engine.py', 'Parsing vykhodnykh faylov FreeSurfer. Izvlekaet obyomy '
     '(Volume) iz aseg.stats i tolshchinu kory (Thickness) iz lh.aparc.stats / '
     'rh.aparc.stats. Podderzhivaet realnyy format faylov s sostavnymi imenami '
     'regionov (Left-Lateral-Ventricle) i avto-poisk v podpapkakh stats/, mri/.'),
    ('qc_core.py', 'Yadro validatsii. Zagruzhaet normy iz ukazannogo istochnika, '
     'vypolnyaet sliyanie (merge) po Region i Type, vychislyaet Z-score, '
     'klassifitsiruet regiony. Podderzhivaet 4 istochnika norm: csv, freesurfer, '
     'enigma, brainchart.'),
    ('load_norms_brainchart.py', 'Zagruzka norm Brain Charts. Chitaet pickle-modeli '
     'norm_blr iz arkhiva Zenodo (57 tys. subektov, 82 ploshchadki). '
     'Realizuet kompozitnyy mapping Destrieux -> Desikan-Killiany dlya tolshchiny '
     '(pooled mean/std po gyrus + sulcus + transitional komponentam).'),
    ('normative_tables.py', 'Dispetcher zagruzki norm iz CSV, FreeSurfer aseg.stats '
     'i ENIGMA (zaglushka).'),
    ('visualizer.py', 'Vizualizatsiya srezov: CLAHE/protsentilnaya normalizatsiya '
     'kontrasta, overlay zalivka + kontur (4-connected numpy eroziya), '
     'krugovye markery dlya kortikalnykh regionov, fallback-ramka. '
     'Podderzhivaet aseg (2-85) i DK-leybly (1000+).'),
    ('input_handler.py', 'Validatsiya vkhodnykh dannykh: proverka nalichiya '
     'obyazatelnykh faylov, BIDS-avtopoisk subekta (3 strategii), fallback-imena '
     'faylov (brain.mgz->T1.mgz, aseg.mgz->aparc+aseg.mgz).'),
    ('app.py', 'Glavnoe Streamlit-prilozhenie: dvukhkolonochnyy dashbord, '
     'interaktivnyy vyuver srezov, selektory istochnika norm, chekboksy '
     'regionov, navigatsiya, eksport CSV/PNG, kastomizatsiya stiley.'),
]

for name, desc in modules:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(' — ' + desc).font.size = Pt(11)

doc.add_heading('2.2. Parsing i normalizatsiya dannykh', level=2)
doc.add_paragraph(
    'Modul parser_engine.py chitaet aseg.stats (podkorkovye obyomy) i '
    'lh/rh.aparc.stats (kortikalnaya tolshchina). Dlya aseg.stats parser '
    'ishchet stroku "# ColHeaders" i izvlekaet StructName (kolonka 4) i '
    'Volume_mm3 (kolonka 3). Sostavnye imena (Left-Lateral-Ventricle) '
    'korrektno obrabatyvayutsya cherez skanirovanie float-tokenov. '
    'Dlya aparc.stats izvlekaetsya StructName i ThickAvg (kolonka 4), '
    'dobavlyaetsya prefiks polushariya (lh-/rh-). '
    'Rezultat — DataFrame s kolonkami Region, Value, Type.'
)

doc.add_heading('2.3. Dva istochnika normativnykh dannykh', level=2)
doc.add_paragraph(
    'Bazovye normy (norms.csv). CSV-fayl s 72 zapisyami: 36 regionov obyoma '
    '(aseg) i 36 regionov tolshchiny (aparc). Znacheniya Mean i Std polucheny '
    'iz literatury i skorrektirovany dlya demonstratsionnykh tseley. '
    'Nedostatok: okhvatyvaet tolko 38 iz 68 aparc-regionov tolshchiny, '
    'chto dayot 39 "Unknown" pri polnom progone QC.'
)
doc.add_paragraph(
    'Nauchnye normy Brain Charts (Nature 2022). Dannye Rutherford et al. '
    '"Charting brain growth and aging at high spatial precision" (eLife, 2022), '
    '57 000+ subektov, 82 issledovatelskie ploshchadki. Dannye skachany s Zenodo '
    '(record 5535467) i raspakovany v brainchart_norms/. '
    'Normy khranyatsya kak pickle-modeli norm_blr; iz meta_data.md izvlekayutsya '
    'mean_resp i std_resp dlya kazhdogo regiona.'
)

doc.add_paragraph(
    'Problema atlasov. FreeSurfer aparc.stats ispolzuet atlas Desikan-Killiany '
    '(34 regiona na polusharie), a Brain Charts — atlas Destrieux '
    '(74+ regionov: otdelno gyrus, sulcus, perekhodnaya zona). '
    'Pryamoy 1:1 mapping privodit k sistematicheskomu sdvigu Z-score '
    '(vse regiony vyglyadyat "khuzhe normy", sredniy Z = -2.0). '
    'Reshenie: kompozitnyy mapping — kazhdyy DK-region = pooled mean/std '
    'po vsem ego Destrieux-komponentam (G_* + S_* + G&S_*) s oboikh polushariy. '
    'Pooled std = sqrt(avg(sigma^2) + var(mu)), chto korrektno uchityvaet i '
    'vnutrikomponentnuyu, i mezhkomponentnuyu dispersiyu. '
    'Rezultat: 81.4% regionov klassifitsiruyutsya kak OK, chto sootvetstvuet '
    'ozhidaemomu raspredeleniyu dlya zdorovogo subekta.'
)

doc.add_heading('2.4. Interaktivnyy dashbord', level=2)
doc.add_paragraph(
    'Streamlit-prilozhenie (app.py) predostavlyaet dvukhkolonochnyy interfeys:'
)

features = [
    'Levaya kolonka: spisok problemnykh regionov s tsvetovoy indikatsiey '
    '(krasnyy — Bad, oranzhevyy — Check), selectbox dlya navigatsii, '
    'knopka "Pereyti k regionu", chekboksy dlya vklyucheniya/vyklyucheniya '
    'kazhdogo regiona, diagnostika faylov.',
    'Pravaya kolonka: vybor istochnika norm (radio), grafiki (gistogramma '
    'Z-score + stolbchataya diagramma verdiktov), tablitsy rezultatov, '
    'interaktivnyy vyuver srezov.',
    'Vyuver: radio-pereklyuchatel osi (Axial/Coronal/Sagittal), '
    'slayder nomera sreza, alpha-polzunki dlya Bad i Check, '
    'chekboks testovoy zelenoy zalivki (Left-Hippocampus, SegId 17).',
    'Navigatsiya: best_slice_for_region() perebiraet vse srezy i nakhodit '
    'indeks s maksimalnoy ploshchadyu maski — knopka "Pereyti" vedyot '
    'na optimalnyy srez, a ne prosto na tsentr mass.',
    'Eksport: st.download_button dlya CSV (qc_results.csv) i PNG (tekushchiy srez).',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.5. Vizualizatsiya srezov (overlay)', level=2)
doc.add_paragraph('Modul visualizer.py realizuet:')
viz_features = [
    'CLAHE-normalizatsiyu kontrasta (scikit-image), s protsentilnym fallback (p2-p98).',
    'Overlay-zalivku: build_overlay_rgba() sozdayot RGBA-sloy s reguliruemoy '
    'prozrachnostyu alpha (0.35 dlya Check, 0.4 dlya Bad).',
    'Kontur: build_contour_rgba() — 4-connected numpy-eroziya bez scipy, '
    'yarkaya granitsa s alpha=1.0.',
    'Markery dlya kortikalnykh regionov: krugi + podpisi, kogda maska '
    'ne vidna na tekushchem sreze (ili kogda aparc+aseg.mgz nedostupen).',
    'Fallback-ramka: shtrikhovaya ramka po tsentru sreza s perechisleniem regionov, '
    'dlya kotorykh net aseg-pokrytiya.',
    'Sinteticheskiy aparc+aseg.mgz: 61 DK-metka, narisovannaya sferami '
    'v anatomicheskikh tsentrakh poverkh aseg.mgz bez zatiraniya podkorkovykh struktur. '
    'Obespechivaet polnotsennuyu zalivku dlya vsekh flagged-regionov.',
]
for f in viz_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.6. Keshirovanie i proizvoditelnost', level=2)
doc.add_paragraph(
    'Dlya uskoreniya pereklyucheniya srezov ispolzuetsya kesh masok v vide dict '
    '(klyuch: label_id + axis + slice_idx). Maski vychislyayutsya odin raz pri '
    'pervom obrashchenii k srezu i pereispolzuyutsya pri povtornykh prokhodakh. '
    'Kesh ochishchaetsya pri novom zapuske QC. Dopolnitelno st.cache_data '
    'primenen dlya zagruzki brain.mgz (tyazhelyy I/O) i aseg-faylov.'
)

doc.add_page_break()

# ============================================================
# 3. SCREENSHOTS
# ============================================================
doc.add_heading('3. Demonstratsiya raboty prototipa', level=1)

screenshots = [
    ('screenshots/1.png', 'Risunok 1 — Glavnyy ekran posle QC: dvukhkolonochnyy '
     'dashbord, sleva — spisok problemnykh regionov, sprava — grafiki i tablitsy.'),
    ('screenshots/2.png', 'Risunok 2 — Levaya kolonka: selectbox navigatsii, '
     'knopka "Pereyti k regionu", chekboksy flagged-regionov, '
     'tsvetovaya indikatsiya verdiktov.'),
    ('screenshots/3.png', 'Risunok 3 — Interaktivnyy vyuver sreza: overlay-zalivka '
     'problemnogo regiona (krasnyy — Bad, oranzhevyy — Check), '
     'kontur i marker s podpisyu.'),
    ('screenshots/4.png', 'Risunok 4 — Pereklyuchatel istochnika normativnykh dannykh: '
     'CSV (norms.csv) i Brain Charts (Nature 2022).'),
    ('screenshots/5.png', 'Risunok 5 — Knopki eksporta rezultatov: '
     'CSV (qc_results.csv) i PNG (tekushchiy srez).'),
    ('screenshots/6.1.png', 'Risunok 6.1 — Navigatsiya "Pereyti k regionu" '
     'do nazhatiya knopki.'),
    ('screenshots/6.2.png', 'Risunok 6.2 — Posle nazhatiya: slayder avtomaticheski '
     'peremeshchen na optimalnyy srez (best_slice_for_region), '
     'otobrazhaetsya overlay problemnogo regiona.'),
]

for rel_path, caption in screenshots:
    full_path = os.path.join(proj, rel_path)
    if os.path.exists(full_path):
        doc.add_picture(full_path, width=Cm(14))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.italic = True
        doc.add_paragraph('')
    else:
        doc.add_paragraph('[Izobrazhenie ne naydeno: %s]' % rel_path)

doc.add_page_break()

# ============================================================
# 4. TABLES AND GRAPHS
# ============================================================
doc.add_heading('4. Tablitsy i grafiki', level=1)

# --- Table 1 ---
doc.add_heading('Tablitsa 1 — Rezultaty QC po regionam (Brain Charts)', level=2)
doc.add_paragraph(
    'V tablitse privedeny pervye 12 strok rezultatov QC pri ispolzovanii '
    'norm Brain Charts (Nature 2022). Dannye polucheny na testovom subekte '
    'sub-01. Z-score vychislen kak (Value - Mean) / Std. '
    'Verdikt: OK — |Z| < 2, Check — 2 <= |Z| < 3, Bad — |Z| >= 3.'
)

table1 = doc.add_table(rows=13, cols=7)
table1.style = 'Light Grid Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ['Region', 'Type', 'Value', 'Mean', 'Std', 'Zscore', 'Verdict']
qc_data = [
    ['Left-Lateral-Ventricle', 'Volume', '12 121.7', '11 009.50', '7 561.38', '0.15', 'OK'],
    ['Left-Cerebellum-Cortex', 'Volume', '55 017.6', '55 576.79', '6 398.60', '-0.09', 'OK'],
    ['Left-Thalamus-Proper', 'Volume', '8 496.0', '7 373.83', '1 114.85', '1.01', 'OK'],
    ['Left-Putamen', 'Volume', '5 280.1', '4 841.57', '715.79', '0.61', 'OK'],
    ['Left-Hippocampus', 'Volume', '3 666.7', '3 955.72', '426.10', '-0.68', 'OK'],
    ['Left-Amygdala', 'Volume', '1 921.7', '1 543.49', '225.97', '1.67', 'OK'],
    ['Right-Lateral-Ventricle', 'Volume', '9 378.9', '10 010.15', '6 721.65', '-0.09', 'OK'],
    ['Right-Hippocampus', 'Volume', '4 258.7', '4 106.70', '447.16', '0.34', 'OK'],
    ['lh-bankssts', 'Thickness', '2.187', '2.679', '0.170', '-2.90', 'Check'],
    ['lh-caudalanteriorcingulate', 'Thickness', '2.247', '2.892', '0.191', '-3.37', 'Bad'],
    ['rh-cuneus', 'Thickness', '1.639', '2.168', '0.257', '-2.06', 'Check'],
    ['rh-pericalcarine', 'Thickness', '1.380', '2.038', '0.214', '-3.07', 'Bad'],
]

for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)

for row_idx, row_data in enumerate(qc_data):
    for col_idx, val in enumerate(row_data):
        cell = table1.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph('Istochnik: rezultaty raboty prototipa, testovyy subekt sub-01, '
                   'normy Brain Charts (Nature 2022).')

# --- Table 2 ---
doc.add_heading('Tablitsa 2 — Sravnenie istochnikov norm (svodka verdiktov)', level=2)
doc.add_paragraph(
    'Sravnenie raspredeleniya verdiktov dlya odnogo i togo zhe subekta '
    '(TEST DATA/sub-01, 113 regionov: 45 obyomov + 68 tolshchiny) '
    'pri ispolzovanii dvukh istochnikov normativnykh dannykh.'
)

table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Verdikt', 'CSV (norms.csv)', 'Brain Charts (Nature 2022)']
data2 = [
    ['OK', '73 (64.6%)', '92 (81.4%)'],
    ['Check', '1 (0.9%)', '6 (5.3%)'],
    ['Bad', '0 (0.0%)', '2 (1.8%)'],
    ['Unknown', '39 (34.5%)', '13 (11.5%)'],
]

for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

for row_idx, row_data in enumerate(data2):
    for col_idx, val in enumerate(row_data):
        cell = table2.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph('')
doc.add_paragraph(
    'Istochnik: rezultaty raboty prototipa. Preimushchestvo Brain Charts '
    'obuslovleno polnym pokrytiem aparc-regionov (68 iz 68 vs 38 iz 68 u CSV) '
    'i ispolzovaniem kompozitnogo DK-mappinga. 13 Unknown u Brain Charts — '
    'struktury bez norm v Brain Charts: 5th-Ventricle, mozolistoe telo (CC_*), '
    'WM-hypointensities, Optic-Chiasm.'
)

doc.add_paragraph('')

# --- Graphs ---
doc.add_heading('4.1. Grafiki', level=2)

graph7_path = os.path.join(proj, 'screenshots/7.png')
if os.path.exists(graph7_path):
    doc.add_picture(graph7_path, width=Cm(14))
    p = doc.add_paragraph(
        'Grafik 1 — Gistogramma raspredeleniya Z-score. '
        'Os X: znachenie Z-score. Os Y: kolichestvo regionov. '
        'Punktirnye linii: oranzhevye — |Z| = 2 (granitsa OK/Check), '
        'krasnye — |Z| = 3 (granitsa Check/Bad). '
        'Raspredelenie blizko k normalnomu s tsentrom okolo 0, '
        'chto ozhidaemo dlya zdorovogo subekta.'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.italic = True
    doc.add_paragraph('')

graph8_path = os.path.join(proj, 'screenshots/8.png')
if os.path.exists(graph8_path):
    doc.add_picture(graph8_path, width=Cm(12))
    p = doc.add_paragraph(
        'Grafik 2 — Stolbchataya diagramma raspredeleniya verdiktov '
        '(istochnik norm: Brain Charts). Os X: kategoriya verdikta. '
        'Os Y: kolichestvo regionov. Tsveta: zelenyy — OK, '
        'oranzhevyy — Check, krasnyy — Bad. '
        '81.4% regionov klassifitsirovany kak OK.'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.italic = True

doc.add_page_break()

# ============================================================
# 5. ANALYSIS
# ============================================================
doc.add_heading('5. Analiz rezultatov', level=1)

doc.add_heading('5.1. Sootvetstvie postavlennym zadacham', level=2)
doc.add_paragraph(
    'Tsel raboty — sozdanie prototipa avtomaticheskogo QC dlya rezultatov '
    'FreeSurfer — dostignuta. Prototip uspeshno:'
)
analysis_items = [
    'parsit realnye vykhodnye fayly FreeSurfer (aseg.stats, lh/rh.aparc.stats) '
    'i izvlekaet 113 izmereniy (45 obyomov + 68 tolshchiny) dlya odnogo subekta;',
    'sravnivaet izmereniya s normami iz dvukh nezavisimykh istochnikov;',
    'vychislyaet Z-score i klassifitsiruet regiony;',
    'otobrazhaet interaktivnyy dashbord s vizualizatsiey problemnykh zon na srezakh;',
    'obespechivaet eksport rezultatov v CSV i PNG.',
]
for item in analysis_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2. Chto rabotaet', level=2)
working_items = [
    'Parsing: korrektno obrabatyvaet sostavnye imena regionov, BIDS-vlozhennost, '
    'fallback-imena faylov (brain.mgz->T1.mgz, aseg.mgz->aparc+aseg.mgz). '
    'Avtopoisk subekta po trem strategiyam.',
    'Z-score i klassifikatsiya: formula (Value - Mean) / Std vychislyaetsya korrektno, '
    'NaN dlya regionov bez norm, porogi |Z|=2 i |Z|=3 dlya verdiktov.',
    'Dva istochnika norm: CSV — prostoe pereklyuchenie, Brain Charts — '
    'nauchno obosnovannye kompozitnye normy s pokrytiem 100% aparc-regionov.',
    'Interaktivnaya vizualizatsiya: overlay s zalivkoy i konturom dlya aseg-regionov, '
    'markery i podpisi dlya DK-regionov, testovaya zelenaya zalivka gippokampa.',
    'Navigatsiya: best_slice_for_region() nakhodit srez s maksimalnoy ploshchadyu '
    'maski — relevantnee, chem tsentr mass.',
    'Chekboksy: polzovatel mozhet vremenno skryvat meshayushchie regiony.',
    'Eksport: CSV-fayl s polnymi rezultatami QC, PNG-snimok tekushchego sreza.',
]
for item in working_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3. Chto trebuet dorabotki', level=2)
todo_items = [
    'Proizvoditelnost: pereklyuchenie srezov vyzyvaet pereschyot masok. '
    'Tekushchiy kesh (dict) smyagchaet problemu, no dlya production nuzhen '
    'predvychislennyy kesh vsekh masok pri zagruzke subekta.',
    'aparc+aseg.mgz: sinteticheskiy fayl s 61 DK-metkoy rabotaet, no dlya '
    'realnogo ispolzovaniya nuzhen libo zapusk recon-all, libo zagruzka '
    'gotovogo fayla iz payplayna.',
    'Vozrastnye normy: Brain Charts predostavlyaet model, zavisyashchuyu ot vozrasta, '
    'no tekushchaya versiya ispolzuet tolko obshchee srednee (NM_0_0_estimate.pkl). '
    'Dlya klinicheskogo primeneniya nuzhen uchyot vozrasta subekta cherez GAM-modeli.',
    'Rasshirenie normativnoy bazy: dobavlenie ENIGMA-norm (trebuet dostupa '
    'k dannym konsortsiuma), norm ploshchadi poverkhnosti.',
    'UI/UX: bazovaya kastomizatsiya Streamlit cherez CSS uzhe primenena '
    '(skruglennye knopki, tsvetovye bedzhi), no dlya produktovogo vida '
    'trebuetsya bolee glubokaya kastomizatsiya.',
]
for item in todo_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 6. CONCLUSION
# ============================================================
doc.add_heading('6. Zaklyuchenie', level=1)
doc.add_paragraph(
    'Razrabotannyy prototip demonstriruet rabotosposobnost podkhoda '
    'k avtomaticheskomu kontrolyu kachestva segmentatsii FreeSurfer na osnove '
    'sravneniya s populyatsionnymi normami. Prototip uspeshno integriruet '
    'dva istochnika normativnykh dannykh, vklyuchaya nauchnye normy Brain Charts '
    '(Nature 2022) s kompozitnym mappingom atlasov, i predostavlyaet '
    'interaktivnyy dashbord dlya vizualnogo analiza problemnykh regionov.'
)
doc.add_paragraph(
    'Prototip gotov k demonstratsii i mozhet sluzhit osnovoy dlya dalneyshey '
    'razrabotki polnotsennogo instrumenta QC. Klyuchevye napravleniya razvitiya: '
    'uchyot vozrasta subekta, predvychislenie masok dlya uskoreniya, '
    'integratsiya s realnymi payplaynami obrabotki (XNat, CBRAIN), '
    'dobavlenie norm ENIGMA i ploshchadi poverkhnosti.'
)

doc.add_paragraph('')

# ============================================================
# 7. APPENDIX
# ============================================================
doc.add_heading('7. Prilozhenie', level=1)
doc.add_paragraph('Repozitoriy s iskhodnym kodom: [SSYLKA NA GITHUB]')
doc.add_paragraph(
    'Sostav repozitoriya: app.py (glavnoe prilozhenie), parser_engine.py, '
    'qc_core.py, load_norms_brainchart.py, normative_tables.py, '
    'visualizer.py, input_handler.py, generate_fake_data.py, '
    'norms.csv, requirements.txt, README.md, report.md.'
)
doc.add_paragraph(
    'Dlya zapuska: pip install -r requirements.txt && '
    'streamlit run app.py'
)
doc.add_paragraph(
    'Testovye dannye: TEST DATA/sub-01/ — rezultaty obrabotki FreeSurfer '
    '(brain.mgz, aseg.mgz, aseg.stats, lh.aparc.stats, rh.aparc.stats, '
    'sinteticheskiy aparc+aseg.mgz).'
)

# --- Save ---
doc.save(out_path)
print('Saved: %s' % out_path)
print('Size: %d bytes' % os.path.getsize(out_path))
print('Done.')
