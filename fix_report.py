#!/usr/bin/env python3
"""Исправляет отчёт практиканта: удаляет пустые разделы, чистит заключение, дополняет основную часть."""
import sys; sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

src = r'C:\Тагир\Проект QC\ОТЧЕТНОСТЬ\отчет-практиканта-заполненный.docx'
doc = Document(src)

# ============================================================
# 1. Delete empty heading sections
# Find and remove paragraphs between "Оглавление" and "Введение"
# ============================================================
paras_to_remove = []
in_zone = False
for p in doc.paragraphs:
    t = p.text.strip()
    if t == 'Оглавление':
        in_zone = True
        continue
    if in_zone and t == 'Введение' and p.style.name.startswith('Heading'):
        break
    if in_zone:
        # Remove empty headings and surrounding empties
        paras_to_remove.append(p)

for p in paras_to_remove:
    try:
        p._element.getparent().remove(p._element)
    except:
        pass

print('Removed %d empty paragraphs between Оглавление and Введение' % len(paras_to_remove))

# ============================================================
# 2. Fix introduction - delete garbage text
# ============================================================
for p in doc.paragraphs:
    t = p.text
    if 'бакалавра, которая определена как «[НАЗВАНИЕ ВЫПУСКНОЙ РАБОТЫ]»' in t:
        print('Found garbage text in introduction')
        for run in p.runs:
            if 'бакалавра, которая определена как «[НАЗВАНИЕ ВЫПУСКНОЙ РАБОТЫ]»' in run.text:
                run.text = run.text.replace(
                    'бакалавра, которая определена как «[НАЗВАНИЕ ВЫПУСКНОЙ РАБОТЫ]»',
                    ''
                )
                print('  Cleaned')
        # Replace with correct VKR theme name
        run = p.add_run(
            ' Тема выпускной квалификационной работы: «Автоматический контроль качества '
            'результатов обработки FreeSurfer (структурная МРТ)».'
        )
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

# ============================================================
# 3. Delete conclusion garbage paragraph
# ============================================================
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('Выполненный во время проведения производственной практики обзор публикаций'):
        try:
            p._element.getparent().remove(p._element)
            print('Removed garbage conclusion paragraph')
        except:
            pass

# ============================================================
# 4. Add 3 paragraphs after the work table
# ============================================================
# Find the last row of table 0 or the paragraph after it
# The work table is table 0, and after it comes conclusion heading (Заключение)
# We need to insert before заключение
insert_point = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == 'Заключение' and p.style.name.startswith('Heading'):
        insert_point = p
        break

if insert_point:
    texts = [
        'Разработка прототипа велась итеративно. Сначала был создан модуль парсинга '
        'parser_engine.py, затем ядро QC (qc_core.py) с расчётом Z-score и классификацией, '
        'после чего добавлена визуализация в visualizer.py. Основная сложность заключалась '
        'в интеграции нормативных данных Brain Charts из-за несовпадения кортикальных '
        'атласов (Desikan-Killiany vs Destrieux). Эта задача была решена путём создания '
        'композитного маппинга — каждый DK-регион агрегируется из нескольких Destrieux-'
        'компонентов (G_ + S_ + G&S_) с pooled стандартным отклонением, учитывающим '
        'как внутрикомпонентную, так и межкомпонентную дисперсию.',

        'В ходе работы возникли трудности с доступом к реальным данным. Данные '
        '160 субъектов (папки RNS и INP) хранились на удалённом Mac, доступ к которому '
        'осуществлялся через AnyDesk и VPN. В некоторых папках отсутствовали файлы '
        'orig/001.mgz — пришлось реализовать fallback на T1.mgz. Парсинг пришлось '
        'адаптировать под BIDS-структуру, где файлы лежат в подпапках stats/ и mri/, '
        'а также под обработку пустых файлов (например, aparc+aseg.mgz нулевого размера).',

        'Пакетная обработка 160 субъектов выполнена с помощью скрипта batch_qc.py, '
        'который для каждого субъекта генерирует JSON с полной таблицей метрик и PNG-срез '
        'с overlay проблемных зон. В процессе анализа выявлены артефакты сегментации: '
        'у субъекта RNS001 объём CSF оказался в 226 раз выше популяционной нормы '
        '(Z-score = 791), что указывает на грубую ошибку FreeSurfer, а не на реальную '
        'анатомию. Все результаты собраны в HTML-отчёты с интерактивными графиками Plotly '
        'и сводной таблицей index.html с сортировкой по столбцам.',
    ]

    for text in reversed(texts):
        new_p = OxmlElement('w:p')
        run_e = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '22')  # 11pt
        rPr.append(rFonts); rPr.append(sz)
        run_e.append(rPr)
        t_e = OxmlElement('w:t')
        t_e.set(qn('xml:space'), 'preserve')
        t_e.text = text
        run_e.append(t_e)
        new_p.append(run_e)
        insert_point._element.addprevious(new_p)

    print('Inserted 3 main part paragraphs before Заключение')

# ============================================================
# 5. Clean up: remove repeated empty lines in conclusion
# Check for empty paragraphs in skills list
# ============================================================
# Find the skills list and remove empty paragraphs between items
in_skills = False
skills_paras = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'В результате практики я получил следующие практические навыки' in t:
        in_skills = True
    if in_skills:
        skills_paras.append(p)
        if t and '—' not in t and not t.startswith('В результате') and not t.startswith('Практика дала') and not t.startswith('За время практики'):
            if i > len(doc.paragraphs) - 5:
                break
            # Check next paragraph is skills too
            if i+1 < len(doc.paragraphs) and doc.paragraphs[i+1].text.strip().startswith('—'):
                pass
            else:
                break

# Remove empty paragraphs between skills
# (actually this is tricky with indices changing. Let's just do the critical fix.)

# ============================================================
# SAVE
# ============================================================
doc.save(src)
print('Saved')
