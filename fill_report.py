#!/usr/bin/env python3
"""Заполняет отчёт практиканта на основе шаблона."""
import sys; sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt

template = r'C:\Тагир\Проект QC\ОТЧЕТНОСТЬ\Шаблоны Документов\отчет-практиканта-шаблон.docx'
output = r'C:\Тагир\Проект QC\ОТЧЕТНОСТЬ\отчет-практиканта-заполненный.docx'

doc = Document(template)

# ============================================================
# TITLE PAGE FIXES
# ============================================================
for p in doc.paragraphs:
    text = p.text

    # Fix student name
    if '____________________' in text and p.style.name == 'Normal':
        for run in p.runs:
            if '____________________' in run.text and 'Ф.И.О' not in run.text:
                run.text = run.text.replace('____________________', 'Гусейнов Тагир Гамзатович')

    # Fix group
    if 'НБИбд-00-23' in text:
        for run in p.runs:
            if 'НБИбд-00-23' in run.text:
                run.text = run.text.replace('НБИбд-00-23', 'НБИбд-01-23')

    # Remove Internet of Things references
    if 'Интернет вещей' in text:
        for run in p.runs:
            if 'Интернет вещей' in run.text:
                run.text = run.text.replace('Интернет вещей', 'QC FreeSurfer')
    if 'характеристик каналов' in text:
        for run in p.runs:
            run.text = run.text.replace('характеристик каналов', 'автоматического контроля качества FreeSurfer')
    if 'дальнейшего прогнозирования' in text:
        for run in p.runs:
            run.text = run.text.replace('дальнейшего прогнозирования', 'формирования отчётов')

# ============================================================
# TABLE 0: Work plan — replace ALL rows
# ============================================================
t0 = doc.tables[0]
# Keep header row (0), replace rows 1-9
work_plan = [
    ('', 'Установочное занятие. Инструктаж.',
     'Обсуждение темы ВКР, постановка задачи: разработка прототипа автоматического QC FreeSurfer.',
     '07.02.2026'),
    ('', 'Изучение литературы и инструментов QC.',
     'Анализ форматов файлов FreeSurfer. Определение стека технологий (Python, NiBabel, Pandas, Streamlit).',
     '10.02–20.02.2026'),
    ('', 'Разработка парсера и ядра QC.',
     'Написан parser_engine.py для извлечения объёмов и толщины коры. Реализован расчёт Z-score и классификация.',
     '21.02–10.03.2026'),
    ('', 'Интеграция норм Brain Charts (Nature 2022).',
     'Загружены данные с Zenodo, создан маппинг Destrieux → Desikan-Killiany. Подключены к QC-ядру.',
     '11.03–20.03.2026'),
    ('', 'Создание веб-прототипа на Streamlit.',
     'Разработан app.py: таблицы, гистограммы, вьювер срезов, overlay, навигация, экспорт CSV/PNG.',
     '21.03–15.04.2026'),
    ('', 'Тестирование на открытых данных.',
     'Проверка парсера, визуализации, экспорта. Исправлены ошибки маппинга кортикальных регионов.',
     '16.04–30.04.2026'),
    ('', 'Разработка пакетной обработки batch_qc.py.',
     'Добавлен расчёт SNR, генерация JSON и PNG для каждого субъекта. Поддержка --single.',
     '01.05–10.05.2026'),
    ('', 'Обработка реальных данных (160 субъектов).',
     'Запуск batch_qc.py на удалённом Mac. Получены JSON и PNG, выявлены артефакты сегментации.',
     '11.05–25.05.2026'),
    ('', 'Генерация HTML-отчётов.',
     'Созданы страницы для каждого субъекта с таблицами, графиками Plotly, встроенными срезами. Сформирован index.html.',
     '26.05–02.06.2026'),
    ('', 'Подготовка отчётных документов.',
     'Заполнение дневника, отзыва, итогового отчёта. Сдача документов.',
     '03.06–07.06.2026'),
]

# Fill work plan table
for i, (num, title, desc, dates) in enumerate(work_plan):
    ri = i + 1  # row index (skip header)
    if ri >= len(t0.rows):
        break
    texts = [str(i+1), title, desc, dates]
    for ci, text in enumerate(texts):
        if ci < len(t0.rows[ri].cells):
            cell = t0.rows[ri].cells[ci]
            cell.text = ''
            cell.paragraphs[0].add_run(text).font.size = Pt(11)
        cell.text = ''
        cell.paragraphs[0].add_run(text).font.size = Pt(11)

# ============================================================
# TABLE 1: Evaluation — mark the safety checkbox
# ============================================================
t1 = doc.tables[1]
# Row 1 (index 1) = instruction checkbox
t1.rows[1].cells[2].text = ''
p = t1.rows[1].cells[2].paragraphs[0]
run = p.add_run('выполнено до 07.02.2026')
run.font.size = Pt(11)

# ============================================================
# INTRODUCTION - rewrite
# ============================================================
# Find the introduction paragraphs and replace
in_intro = False
intro_paras = []
for i, p in enumerate(doc.paragraphs):
    if 'Согласно программе' in p.text and 'целями практики являются' in p.text:
        in_intro = True
    if in_intro and 'Время проведения' in p.text:
        intro_paras.append(i)
    if in_intro and 'Последовательность прохождения' in p.text:
        intro_paras.append(i)
        break

# We'll rewrite sequentially
for p in doc.paragraphs:
    if 'Согласно программе' in p.text and 'целями практики являются' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('Целью учебной практики было получение навыков научно-исследовательской работы '
                          'и разработка прототипа автоматического контроля качества сегментации МРТ-снимков, '
                          'обработанных в программе FreeSurfer. Тема выпускной квалификационной работы: '
                          '«Автоматический контроль качества результатов обработки FreeSurfer (структурная МРТ)».')

    if 'Для достижения целей' in p.text and 'обзор публикаций' in p.text:
        for run in p.runs:
            if 'Для достижения целей' in run.text:
                run.text = ''
        p.runs[0].text = ('За время практики я изучил существующие инструменты контроля качества — '
                          'FSQC, VisualQC, MRIQC — и разработал собственный прототип на Python, '
                          'который парсит файлы FreeSurfer, вычисляет Z-score с использованием '
                          'научных норм Brain Charts (Nature 2022) и показывает проблемные зоны '
                          'на срезах мозга.')

    if 'Время проведения производственной практики' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('Время проведения практики: 07.02.2026 – 13.06.2026. '
                          'Место проведения: Научный центр вычислительных методов и моделирования '
                          'киберфизических систем РУДН.')

    if 'Последовательность прохождения практики, перечень работ' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('Последовательность работ приведена в таблице раздела «Основная часть». '
                          'Все этапы выполнены в установленные сроки.')

    # Fix practice type in heading
    if 'Учебной практики' in p.text and 'Научно-исследовательская' in p.text:
        for run in p.runs:
            if 'Учебной практики' in run.text:
                run.text = run.text.replace('Учебной практики', 'учебной практики')

    # Remove IoT references in conclusion
    if 'обзор публикаций научных изданий как по теме Интернет вещей' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('Выполненный во время практики обзор публикаций и разработанный программный '
                          'прототип показали, что задача автоматического контроля качества сегментации '
                          'FreeSurfer актуальна и может быть решена с помощью описанного подхода.')

# ============================================================
# CONCLUSIONS - rewrite
# ============================================================
for p in doc.paragraphs:
    if 'За период практики, которая осуществлялась' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('За время практики я освоил библиотеки для работы с нейровизуализационными данными '
                          '(NiBabel, Pandas, NumPy), научился делать интерактивные веб-интерфейсы на Streamlit, '
                          'интегрировать научные данные из открытых архивов (Zenodo) и обрабатывать большие '
                          'массивы данных (160 субъектов, ~18 тысяч метрик).')

    if 'При прохождении практики я разобрался с научной терминологией' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('Практика дала понимание того, как устроены нейровизуализационные исследования, '
                          'какие форматы данных используются (NIfTI, .mgz, .stats), как работают атласы '
                          'головного мозга (Desikan-Killiany, Destrieux) и как можно автоматизировать '
                          'проверку качества на уровне отдельных анатомических структур.')

    if 'В результате прохождения данной практики мной были получены' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('В результате практики я получил следующие практические навыки:')

    # Replace generic competence list with shorter relevant ones
    if 'способностью к самоорганизации' in p.text and 'к самообразованию' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— способность самостоятельно ставить задачу, выбирать инструменты и доводить работу до результата;'

    if 'способностью решать стандартные задачи' in p.text and 'библиографической культуры' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— умение искать, анализировать и обрабатывать научно-техническую информацию;'

    if 'способностью к самостоятельной научно-исследовательской работе' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— навык написания программ на Python с использованием специализированных библиотек;'

    if 'способностью находить, анализировать, реализовывать программно' in p.text and 'математические алгоритмы' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— умение реализовывать статистические алгоритмы (Z-score) и визуализировать результаты;'

    if 'способностью к определению общих форм и закономерностей' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— понимание структуры и закономерностей нейровизуализационных данных;'

    if 'способностью математически корректно ставить' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— умение формировать отчёты в форматах HTML, JSON, CSV для последующего анализа;'

    if 'способностью строго доказать утверждение' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— навык работы с системами контроля версий (Git) и открытыми данными (OpenNeuro, Zenodo);'

    if 'способностью публично представлять собственные' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— опыт настройки удалённого доступа (AnyDesk, VPN) и работы на macOS;'

    if 'способностью использовать методы математического' in p.text and 'алгоритмического моделирования' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '— навык использования Plotly для интерактивных графиков и Matplotlib для статической визуализации.'

    if 'способностью передавать результат' in p.text and 'конкретных рекомендаций' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ''

    if 'способностью представлять и адаптировать знания' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ''

    if 'способностью к проведению методических и экспертных работ' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ''

# ============================================================
# DELETE LIST OF ABBREVIATIONS (if possible)
# ============================================================
# We can't delete paragraphs easily, but we can clear them
for p in doc.paragraphs:
    if p.text.strip() in ['Список сокращений', 'Список основных обозначений',
                           'Русскоязычные сокращения', 'Англоязычные сокращения']:
        for run in p.runs:
            run.text = ''
    if p.style.name == 'Heading 1' and 'Список сокращений' in p.text:
        for run in p.runs:
            if 'Список сокращений' in run.text:
                run.text = ''

# ============================================================
# REFERENCES — add after "Список источников"
# ============================================================
for i, p in enumerate(doc.paragraphs):
    if 'Список источников' in p.text and p.style.name.startswith('Heading'):
        # Add references after this heading
        # We can add through a paragraph append trick
        refs = [
            '1. Bethlehem R.A.I. et al. Brain charts for the human lifespan // Nature. — 2022. — Vol. 604. — P. 525–533.',
            '2. Fischl B. FreeSurfer // NeuroImage. — 2012. — Vol. 62, № 2. — P. 774–781.',
            '3. Fischl B., Dale A.M. Measuring the thickness of the human cerebral cortex from magnetic resonance images // PNAS. — 2000. — Vol. 97, № 20. — P. 11050–11055.',
            '4. Desikan R.S. et al. An automated labeling system for subdividing the human cerebral cortex on MRI scans // NeuroImage. — 2006. — Vol. 31, № 3. — P. 968–980.',
            '5. Rutherford S. et al. Charting brain growth and aging at high spatial precision // eLife. — 2022. — Vol. 11. — P. e72904.',
            '6. Klapwijk E.T. et al. Qoala-T: A supervised-learning tool for quality control of FreeSurfer segmented MRI data // NeuroImage. — 2019. — Vol. 189. — P. 116–129.',
            '7. Esteban O. et al. MRIQC: Advancing the automatic prediction of image quality // PLoS ONE. — 2017. — Vol. 12, № 9. — P. e0184661.',
            '8. FreeSurfer Documentation [Электронный ресурс]. — URL: https://surfer.nmr.mgh.harvard.edu/fswiki (дата обращения: 01.06.2026).',
            '9. Streamlit Documentation [Электронный ресурс]. — URL: https://docs.streamlit.io (дата обращения: 01.06.2026).',
            '10. OpenNeuro Dataset ds000248 [Электронный ресурс]. — URL: https://openneuro.org/datasets/ds000248 (дата обращения: 01.06.2026).',
        ]
        # Add to the paragraph after the heading
        parent = p._element
        parent.addnext(
            doc.paragraphs[i+1]._element if i+1 < len(doc.paragraphs) else None
        )
        break

# Actually, let's just append references after "Список источников" heading
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Список источников' and p.style.name.startswith('Heading'):
        # Clear any existing content after
        p.text = ''
        p.add_run('Список источников').bold = True
        break

# Check if the list needs to be added - we'll do it in the text by modifying the last paragraph
for i, p in enumerate(doc.paragraphs):
    if 'Список источников' in p.text and 'Heading' not in p.style.name:
        p.clear()
        refs = [
            '1. Bethlehem R.A.I. et al. Brain charts for the human lifespan // Nature. 2022. Vol. 604. P. 525–533.',
            '2. Fischl B. FreeSurfer // NeuroImage. 2012. Vol. 62, № 2. P. 774–781.',
            '3. Desikan R.S. et al. An automated labeling system for subdividing the human cerebral cortex on MRI scans // NeuroImage. 2006. Vol. 31, № 3. P. 968–980.',
            '4. Rutherford S. et al. Charting brain growth and aging at high spatial precision // eLife. 2022. Vol. 11. P. e72904.',
            '5. Klapwijk E.T. et al. Qoala-T: A supervised-learning tool for quality control of FreeSurfer segmented MRI data // NeuroImage. 2019. Vol. 189. P. 116–129.',
            '6. Esteban O. et al. MRIQC: Advancing the automatic prediction of image quality in MRI from unseen sites // PLoS ONE. 2017. Vol. 12, № 9. P. e0184661.',
            '7. FreeSurfer Wiki — Quality Control [Электронный ресурс]. URL: https://surfer.nmr.mgh.harvard.edu/fswiki (дата обращения: 01.06.2026).',
            '8. Streamlit Documentation [Электронный ресурс]. URL: https://docs.streamlit.io (дата обращения: 01.06.2026).',
            '9. OpenNeuro Dataset ds000248 [Электронный ресурс]. URL: https://openneuro.org/datasets/ds000248 (дата обращения: 01.06.2026).',
        ]
        for ref in refs:
            p.add_run('\n' + ref)

# ============================================================
# "Информационно-программные продукты" section
# ============================================================
for p in doc.paragraphs:
    if 'Интернет вещей' in p.text and 'разработки математической модели' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ('Для выполнения работы использовались: Python 3.11, библиотеки pandas, numpy, '
                          'nibabel, matplotlib, streamlit, plotly, tqdm, scikit-image; '
                          'нормативные данные Brain Charts (Zenodo, DOI: 10.7554/eLife.72904); '
                          'система контроля версий Git; удалённый доступ AnyDesk; '
                          'тестовые данные OpenNeuro ds000248.')

# ============================================================
# UPDATE MAIN PART HEADER
# ============================================================
for p in doc.paragraphs:
    if p.style.name.startswith('Heading') and p.text.strip() == 'Основная часть' and 'Heading' in p.style.name:
        break

# Add main part content by modifying the paragraph after "Основная часть"
found_main = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Основная часть' and 'Heading' in p.style.name:
        found_main = True
        continue
    if found_main and p.style.name.startswith('Heading'):
        break
    if found_main and p.text.strip() and 'Heading' not in p.style.name:
        # Replace main content paragraphs
        text = p.text
        if 'Я ознакомился' in text or 'производственной практики' in text:
            for run in p.runs:
                run.text = ''
            p.runs[0].text = ('Я ознакомился с темой исследования, изучил литературу по контролю качества '
                              'в нейровизуализации. Проанализировал существующие инструменты: FSQC, VisualQC, '
                              'MRIQC, Qoala-T. Выявил их ограничения — отсутствие регион-специфичной '
                              'классификации и интерактивной визуализации.')

# ============================================================
# SAVE
# ============================================================
doc.save(output)
print('Saved: %s' % output)
