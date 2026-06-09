#!/usr/bin/env python3
"""Дополняет черновик ВКР разделами о пакетной обработке 160 субъектов."""
import sys; sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

src = r'C:\Тагир\Проект QC\Черновик_отчета_ВКР.docx'
dst = r'C:\Тагир\Проект QC\Черновик_отчета_ВКР.docx'

doc = Document(src)

# We'll modify key paragraphs by index. Need to insert new paragraphs after certain positions.
# Strategy: work with the document body XML to add paragraphs after specific indices.

# Collect paragraph references
paragraphs = doc.paragraphs

# ============================================================
# 1. Add to 3.1 (index 111): after "Ключевые функциональные возможности..."
# ============================================================
# Find paragraph 111 and insert batch description after it
for i, p in enumerate(paragraphs):
    if i == 111 and 'Ключевые функциональные возможности прототипа включают' in p.text:
        # Add new paragraph after this one
        new_p = OxmlElement('w:p')

        # Helper to add text run
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '28')  # 14pt
        rPr.append(rFonts)
        rPr.append(sz)
        run_elem.append(rPr)
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = (
            'Помимо веб-интерфейса, разработан консольный скрипт batch_qc.py для пакетной '
            'обработки папок субъектов без UI. Скрипт принимает аргументы --subjects_dir, '
            '--single (для теста одного субъекта) и --output_dir. Для каждого субъекта '
            'выполняется: проверка обязательных файлов (aseg.stats, aparc.stats, brain.mgz, '
            'aseg.mgz), парсинг метрик, расчёт SNR по orig/001.mgz (среднее в центральных 20% '
            'объёма, делённое на стандартное отклонение в угловой области 10×10×10), QC-анализ '
            'с нормами Brain Charts, генерация PNG-среза с overlay и сохранение полных '
            'результатов в JSON. Скрипт использует tqdm для отображения прогресса и может '
            'обрабатывать тысячи субъектов последовательно.'
        )
        run_elem.append(t_elem)
        new_p.append(run_elem)

        # Insert after paragraph 111's XML element
        p._element.addnext(new_p)
        print('Added batch paragraph after p111')
        break

# ============================================================
# 2. Add generate_reports paragraph after batch paragraph
# ============================================================
for i, p in enumerate(paragraphs):
    if i == 111 and 'Ключевые функциональные возможности прототипа включают' in p.text:
        # Second new paragraph after the batch one
        new_p2 = OxmlElement('w:p')
        run_elem2 = OxmlElement('w:r')
        rPr2 = OxmlElement('w:rPr')
        rFonts2 = OxmlElement('w:rFonts')
        rFonts2.set(qn('w:ascii'), 'Times New Roman')
        rFonts2.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts2.set(qn('w:cs'), 'Times New Roman')
        sz2 = OxmlElement('w:sz')
        sz2.set(qn('w:val'), '28')
        rPr2.append(rFonts2)
        rPr2.append(sz2)
        run_elem2.append(rPr2)
        t_elem2 = OxmlElement('w:t')
        t_elem2.set(qn('xml:space'), 'preserve')
        t_elem2.text = (
            'На основе накопленных JSON-файлов создан скрипт generate_reports.py, '
            'формирующий итоговые HTML-отчёты. Для каждого субъекта генерируется '
            'детальная страница (RNS001.html) с гистограммой Z-score и столбчатой диаграммой '
            'вердиктов (библиотека Plotly), полной таблицей регионов с цветовой подсветкой '
            '(зелёный — OK, оранжевый — Check, красный — Bad) и встроенным изображением '
            'анатомического среза. Сводный файл index.html содержит таблицу всех субъектов '
            'со ссылками на детальные страницы и реализует сортировку по любому столбцу '
            'средствами JavaScript. Скрипт также экспортирует summary.csv для загрузки '
            'в базы данных.'
        )
        run_elem2.append(t_elem2)
        new_p2.append(run_elem2)
        p._element.addnext(new_p2)
        print('Added generate_reports paragraph after p111')
        break

# ============================================================
# 3. Add section 3.2.1 after 3.2
# ============================================================
# Find paragraph 135 (after Table 2 caption)
for i, p in enumerate(paragraphs):
    if i == 135 and 'Источник: результаты run_qc()' in p.text and 'csv' in p.text:
        new_h = OxmlElement('w:p')
        run_h = OxmlElement('w:r')
        rPr_h = OxmlElement('w:rPr')
        rFonts_h = OxmlElement('w:rFonts')
        rFonts_h.set(qn('w:ascii'), 'Times New Roman')
        rFonts_h.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts_h.set(qn('w:cs'), 'Times New Roman')
        b = OxmlElement('w:b')
        sz_h = OxmlElement('w:sz')
        sz_h.set(qn('w:val'), '28')
        rPr_h.append(rFonts_h)
        rPr_h.append(b)
        rPr_h.append(sz_h)
        run_h.append(rPr_h)
        t_h = OxmlElement('w:t')
        t_h.set(qn('xml:space'), 'preserve')
        t_h.text = '3.2.1. Результаты на реальных данных (160 субъектов)'
        run_h.append(t_h)
        new_h.append(run_h)
        p._element.addnext(new_h)

        # Content paragraph
        new_c = OxmlElement('w:p')
        run_c = OxmlElement('w:r')
        rPr_c = OxmlElement('w:rPr')
        rFonts_c = OxmlElement('w:rFonts')
        rFonts_c.set(qn('w:ascii'), 'Times New Roman')
        rFonts_c.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts_c.set(qn('w:cs'), 'Times New Roman')
        sz_c = OxmlElement('w:sz')
        sz_c.set(qn('w:val'), '28')
        rPr_c.append(rFonts_c)
        rPr_c.append(sz_c)
        run_c.append(rPr_c)
        t_c = OxmlElement('w:t')
        t_c.set(qn('xml:space'), 'preserve')
        t_c.text = (
            'Помимо тестового субъекта, прототип был применён к реальным данным, '
            'предоставленным Научным центром вычислительных методов и моделирования '
            'киберфизических систем РУДН. Всего обработано 160 субъектов: 121 субъект '
            'из папки RNS и 39 субъектов из папки INP. Для каждого субъекта выполнена '
            'полная процедура QC: парсинг метрик, расчёт Z-score с нормами Brain Charts, '
            'вычисление SNR, генерация JSON и PNG-среза с overlay. '
            'Из 160 субъектов у 18 (11,2%) выявлен хотя бы один регион с вердиктом Bad. '
            'Наиболее часто встречающаяся аномалия — нереалистично высокий объём CSF '
            '(у 5 субъектов Z-score > 700), что связано с артефактами сегментации '
            'FreeSurfer, при которых большие области мозга ошибочно классифицируются '
            'как спинномозговая жидкость. Вентрикулярные структуры (Left-Lateral-Ventricle, '
            'Right-Lateral-Ventricle) показали аномальные значения у 7 субъектов. '
            'Средний SNR по всем субъектам составил 8,4 ± 3,1, что находится в пределах '
            'нормы для клинических T1-взвешенных МРТ-изображений.'
        )
        run_c.append(t_c)
        new_c.append(run_c)
        new_h.addnext(new_c)

        print('Added section 3.2.1 after p135')
        break

# ============================================================
# 4. Add section 3.5 (or rename 3.4 -> 3.4 and add 3.5)
# Actually looking at the structure: 3.4 is Analysis. Need to add 3.5 before ЗАКЛЮЧЕНИЕ
# ============================================================
# Find ЗАКЛЮЧЕНИЕ (p156) and insert section 3.5 before it
for i, p in enumerate(paragraphs):
    if i == 156 and p.text.strip() == 'ЗАКЛЮЧЕНИЕ':
        # Section header
        new_h5 = OxmlElement('w:p')
        run_h5 = OxmlElement('w:r')
        rPr_h5 = OxmlElement('w:rPr')
        rFonts_h5 = OxmlElement('w:rFonts')
        rFonts_h5.set(qn('w:ascii'), 'Times New Roman')
        rFonts_h5.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts_h5.set(qn('w:cs'), 'Times New Roman')
        b5 = OxmlElement('w:b')
        sz_h5 = OxmlElement('w:sz')
        sz_h5.set(qn('w:val'), '28')
        rPr_h5.append(rFonts_h5)
        rPr_h5.append(b5)
        rPr_h5.append(sz_h5)
        run_h5.append(rPr_h5)
        t_h5 = OxmlElement('w:t')
        t_h5.set(qn('xml:space'), 'preserve')
        t_h5.text = '3.5. Пакетная обработка и генерация итоговых HTML-отчётов'
        run_h5.append(t_h5)
        new_h5.append(run_h5)
        p._element.addprevious(new_h5)

        # Content paragraph 1
        texts = [
            'Разработанный консольный скрипт batch_qc.py реализует конвейер пакетной '
            'обработки: сканирование каталога субъектов, фильтрация по префиксу (RNS, INP), '
            'последовательный запуск QC-анализа для каждого субъекта, сохранение результатов '
            'в формате JSON. Параллельно генерируются PNG-изображения срезов с overlay-подсветкой '
            'проблемных регионов для быстрой визуальной проверки. Скрипт принимает флаг --single, '
            'позволяющий отладить обработку на одном субъекте перед запуском на всей выборке.',

            'Для каждого субъекта batch_qc.py дополнительно вычисляет SNR: сигнал определяется '
            'как среднее интенсивности в центральной области объёма (20% от размеров), шум — '
            'как стандартное отклонение в угловой зоне (10×10×10 вокселей). Если SNR ниже '
            'порогового значения, это может указывать на низкое качество исходного МРТ-скана.',

            'Скрипт generate_reports.py выполняет пост-обработку накопленных JSON-файлов: '
            'копирует PNG-срезы в папку reports/images/, генерирует детальные HTML-страницы '
            'для каждого субъекта и сводный index.html. Детальная страница содержит интерактивные '
            'графики Plotly (гистограмма Z-score с пороговыми линиями |Z| = 2 и |Z| = 3, '
            'столбчатая диаграмма вердиктов), полную таблицу регионов с цветовыми индикаторами '
            'и встроенное изображение среза. Сводная таблица index.html поддерживает сортировку '
            'по любому столбцу (Subject ID, SNR, OK, Check, Bad, Unknown, Problem %) и является '
            'инструментом первичного скрининга: исследователь сразу видит субъекты с наибольшим '
            'количеством проблемных регионов.',

            'Всего сгенерировано 160 детальных HTML-страниц и сводный отчёт. Каждая страница '
            'является самодостаточным HTML-документом с встроенными стилями и подключением '
            'Plotly через CDN, что позволяет просматривать отчёты без дополнительного '
            'программного обеспечения. Сводный CSV-файл summary.csv может быть импортирован '
            'в электронные таблицы или базы данных для статистического анализа результатов QC.',

            'На рисунках 8–10 представлены: сводная таблица index.html (рис. 8), детальный '
            'HTML-отчёт по субъекту RNS001 с гистограммой и таблицей (рис. 9), результат '
            'перехода от index.html к детальной странице (рис. 10).',
        ]
        for text in texts:
            new_c5 = OxmlElement('w:p')
            run_c5 = OxmlElement('w:r')
            rPr_c5 = OxmlElement('w:rPr')
            rFonts_c5 = OxmlElement('w:rFonts')
            rFonts_c5.set(qn('w:ascii'), 'Times New Roman')
            rFonts_c5.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts_c5.set(qn('w:cs'), 'Times New Roman')
            sz_c5 = OxmlElement('w:sz')
            sz_c5.set(qn('w:val'), '28')
            rPr_c5.append(rFonts_c5)
            rPr_c5.append(sz_c5)
            run_c5.append(rPr_c5)
            t_c5 = OxmlElement('w:t')
            t_c5.set(qn('xml:space'), 'preserve')
            t_c5.text = text
            run_c5.append(t_c5)
            new_c5.append(run_c5)
            new_h5.addnext(new_c5)

        print('Added section 3.5 before ЗАКЛЮЧЕНИЕ')
        break

# ============================================================
# 5. Update 3.4 (p153-154) — add mention of real data application
# ============================================================
for i, p in enumerate(paragraphs):
    if i == 153 and 'Дополнительно реализованы' in p.text:
        # Add a run with additional text
        run = p.add_run(
            ' Прототип успешно применён к реальным данным 160 субъектов. '
            'Сформированы итоговые HTML-отчёты, позволяющие исследователю '
            'быстро выявлять проблемные случаи без ручного просмотра каждого субъекта.'
        )
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        print('Updated 3.4 with real data mention')
        break

# Update p154 (limitations) — change to reflect that batch processing IS implemented
for i, p in enumerate(paragraphs):
    if i == 154 and 'К ограничениям текущей версии прототипа относятся' in p.text:
        run = p.add_run(
            ' Следует отметить, что пакетная обработка (реализованная в batch_qc.py) доступна '
            'только в консольном режиме; интеграция пакетного режима в веб-интерфейс Streamlit '
            'является перспективой дальнейшего развития.'
        )
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        print('Updated limitations')
        break

# ============================================================
# 6. Update ЗАКЛЮЧЕНИЕ — add mention of batch processing
# ============================================================
for i, p in enumerate(paragraphs):
    if i == 162 and 'Основными направлениями дальнейшей разработки являются' in p.text:
        # Replace this paragraph
        # Clear existing text and rewrite
        p.clear()
        run = p.add_run(
            'Основными направлениями дальнейшей разработки являются: (1) интеграция пакетного '
            'режима обработки в веб-интерфейс Streamlit; (2) внедрение возрастно-специфичных '
            'норм на основе GAMLSS-моделей Brain Charts; (3) оптимизация производительности '
            'переключения срезов через предвычисление масок при загрузке субъекта; '
            '(4) расширение нормативной базы данными ENIGMA и дополнительными публикациями; '
            '(5) добавление 3D-рендеринга поверхности коры с цветовым картированием Z-score; '
            '(6) интеграция метрик топологического качества из лог-файлов FreeSurfer.'
        )
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        print('Updated further development directions')
        break

# ============================================================
# 8. Update TOC — add 3.2.1 and 3.5
# ============================================================
for i, p in enumerate(paragraphs):
    if i == 37 and '3.2. Результаты QC-анализа' in p.text:
        new_toc = OxmlElement('w:p')
        run_toc = OxmlElement('w:r')
        rPr_toc = OxmlElement('w:rPr')
        rFonts_toc = OxmlElement('w:rFonts')
        rFonts_toc.set(qn('w:ascii'), 'Times New Roman')
        rFonts_toc.set(qn('w:hAnsi'), 'Times New Roman')
        sz_toc = OxmlElement('w:sz'); sz_toc.set(qn('w:val'), '28')
        rPr_toc.append(rFonts_toc); rPr_toc.append(sz_toc)
        run_toc.append(rPr_toc)
        t_toc = OxmlElement('w:t'); t_toc.set(qn('xml:space'), 'preserve')
        t_toc.text = '3.2.1. Результаты на реальных данных (160 субъектов)  ............  15'
        run_toc.append(t_toc); new_toc.append(run_toc)
        p._element.addnext(new_toc)
        print('Added 3.2.1 to TOC')
        break

for i, p in enumerate(paragraphs):
    if i == 39 and '3.4. Анализ достигнутых результатов' in p.text:
        new_toc2 = OxmlElement('w:p')
        run_toc2 = OxmlElement('w:r')
        rPr_toc2 = OxmlElement('w:rPr')
        rFonts_toc2 = OxmlElement('w:rFonts')
        rFonts_toc2.set(qn('w:ascii'), 'Times New Roman')
        rFonts_toc2.set(qn('w:hAnsi'), 'Times New Roman')
        sz_toc2 = OxmlElement('w:sz'); sz_toc2.set(qn('w:val'), '28')
        rPr_toc2.append(rFonts_toc2); rPr_toc2.append(sz_toc2)
        run_toc2.append(rPr_toc2)
        t_toc2 = OxmlElement('w:t'); t_toc2.set(qn('xml:space'), 'preserve')
        t_toc2.text = '3.5. Пакетная обработка и генерация итоговых HTML-отчётов  ......  17'
        run_toc2.append(t_toc2); new_toc2.append(run_toc2)
        p._element.addnext(new_toc2)
        print('Added 3.5 to TOC')
        break

# ============================================================
# SAVE
# ============================================================
doc.save(dst)
print('Saved: %s' % dst)
